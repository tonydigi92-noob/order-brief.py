import streamlit as st
from docx import Document
import io
import datetime
import sqlite3
import pandas as pd
import base64
import os

# --- Page Configuration ---
st.set_page_config(layout="wide", page_title="Pixel Resort Briefs", page_icon="👾")

# --- Database Integration (SQLite) ---
def init_db():
    conn = sqlite3.connect('briefs.db', check_same_thread=False)
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS orders (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            order_name TEXT,
            platform TEXT,
            deadline TEXT,
            content_summary TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    conn.commit()
    return conn

conn = init_db()

def save_to_db(order_name, platform, deadline, content_summary):
    c = conn.cursor()
    c.execute('INSERT INTO orders (order_name, platform, deadline, content_summary) VALUES (?, ?, ?, ?)',
              (order_name, platform, deadline, content_summary))
    conn.commit()
    st.toast("💾 Saved to Memory Card!", icon="👾")

def load_history():
    return pd.read_sql_query("SELECT * FROM orders ORDER BY created_at DESC", conn)

# --- Background Image Logic ---
def get_base64(bin_file):
    with open(bin_file, 'rb') as f:
        data = f.read()
    return base64.b64encode(data).decode()

bg_style = "background: linear-gradient(to bottom, #87CEEB, #E0F7FA);" # Fallback
if os.path.exists("bg.jpg"):
    bin_str = get_base64("bg.jpg")
    bg_style = f"""
        background-image: url("data:image/png;base64,{bin_str}");
        background-size: cover;
        background-repeat: no-repeat;
        background-attachment: fixed;
    """

# --- CSS Injection (Retro Resort Theme) ---
st.markdown(f"""
<style>
    /* Google Font for Monospace */
    @import url('https://fonts.googleapis.com/css2?family=VT323&family=Courier+New&display=swap');

    /* Global Settings */
    .stApp {{
        {bg_style} /* <--- BACKGROUND CHANGED HERE */
        font-family: 'Courier New', monospace;
        color: #1A237E;
    }}
    
    /* Remove Top Padding (Adjusted to avoid cutoff) */
    .block-container {{
        padding-top: 4rem !important;
        padding-bottom: 5rem !important;
    }}

    /* Headers */
    h1, h2, h3, h4, .main-header {{
        font-family: 'VT323', 'Courier New', monospace;
        color: #1A237E !important;
        text-transform: uppercase;
    }}
    
    /* Input Fields (Robust Targeting for White Background + Retro Cursor) */
    .stTextInput input, .stTextArea textarea, .stSelectbox > div > div, .stDateInput > div > div, .stTimeInput > div > div, div[data-baseweb="select"] > div, div[data-baseweb="base-input"] {{
        background-color: #FFFFFF !important;
        color: #000000 !important;
        border: 2px solid #1A237E !important;
        border-radius: 0px !important;
        font-family: 'Courier New', monospace !important;
        caret-color: #FF6F00 !important; /* Retro Orange Cursor */
    }}
    
    /* Force Text Color inside Select/Inputs */
    div[data-baseweb="select"] span, div[data-baseweb="base-input"] input {{
        color: #000000 !important;
        -webkit-text-fill-color: #000000 !important;
    }}
    
    /* Labels (Add Background to make them pop) */
    .stMarkdown label, .stTextInput label, .stSelectbox label, .stTextArea label, .stDateInput label, .stTimeInput label {{
        color: #1A237E !important;
        background-color: rgba(255, 255, 255, 0.95); /* High contrast background */
        padding: 4px 12px; /* Increased padding */
        border: 2px solid #1A237E;
        box-shadow: 2px 2px 0px #1A237E;
        margin-bottom: 8px; /* More space below label */
        display: inline-block; 
        font-weight: bold !important;
        font-size: 0.9rem !important;
        border-radius: 0px;
    }}
    
    /* Buttons (Orange Pop) */
    .stButton>button {{
        background-color: #FF6F00 !important;
        color: #FFFFFF !important;
        border: 2px solid #1A237E !important;
        border-radius: 0px !important;
        font-family: 'Courier New', monospace !important;
        font-weight: bold !important;
        text-transform: uppercase;
        box-shadow: 4px 4px 0px #1A237E; /* Pixel shadow */
        transition: all 0.1s;
    }}
    .stButton>button:hover {{
        transform: translate(2px, 2px);
        box-shadow: 2px 2px 0px #1A237E;
    }}
    .stButton>button:active {{
        transform: translate(4px, 4px);
        box-shadow: none;
    }}

    /* Tabs */
    .stTabs [data-baseweb="tab-list"] button {{
        background-color: rgba(255,255,255,0.5);
        border: 2px solid #1A237E;
        color: #1A237E;
        border-radius: 0px;
        margin-right: 5px;
        font-family: 'Courier New', monospace;
        font-weight: bold;
    }}
    .stTabs [aria-selected="true"] {{
        background-color: #FF6F00 !important;
        color: white !important;
    }}

    /* Preview Card */
    .preview-card {{
        background-color: rgba(255, 255, 255, 0.9);
        border: 2px solid #1A237E;
        padding: 1.5rem;
        box-shadow: 8px 8px 0px rgba(26, 35, 126, 0.2);
    }}

    /* Compact Spacing */
    div[data-testid="column"] {{
        gap: 0.5rem;
    }}
    
    /* Helper for 'Custom: size' input to look integrated */
    div[data-testid="stExpander"] {{
        background-color: rgba(255,255,255,0.6);
        border: 1px dashed #1A237E;
        border-radius: 0px;
    }}
</style>
""", unsafe_allow_html=True)

# --- Header ---
# --- Header ---
st.markdown("<h1 style='text-align: center; margin-bottom: 5px;'>🌴 Order Brief WSKL 🌴</h1>", unsafe_allow_html=True)
st.markdown("<p style='text-align: center; font-size: 0.8rem; font-weight: bold; margin-bottom: 20px;'>make by Tony Nguyen</p>", unsafe_allow_html=True)

# --- Tabs ---
tab1, tab2 = st.tabs(["📝 BRIEF BUILDER", "💾 MEMORY CARD"])

with tab1:
    col_input, col_preview = st.columns([1.2, 1], gap="medium")

    # --- LEFT COLUMN: INPUTS ---
    with col_input:
        # 1. Order Grid
        r1c1, r1c2, r1c3 = st.columns([0.5, 0.25, 0.25])
        with r1c1:
            order_name = st.text_input("ORDER NAME", placeholder="QC Tet 2026")
        with r1c2:
            deadline_date = st.date_input("DATE", datetime.date.today())
        with r1c3:
            deadline_time = st.time_input("TIME", datetime.time(17, 0))
            
        deadline_str = f"{deadline_date.strftime('%d/%m/%Y')} {deadline_time.strftime('%H:%M')}"
        
        # 2. Platform Grid
        PLATFORMS = {
            "Facebook": ["Single Image", "Album Grid", "Cover", "Story"],
            "Instagram": ["Post (Square)", "Post (Portrait)", "Story"],
            "Offline": ["Standee", "Poster", "Backdrop"],
            "Digital": ["TV Screen", "Web Banner"],
            "Other": ["Custom"]
        }
        
        r2c1, r2c2 = st.columns([0.4, 0.6])
        with r2c1:
            platform = st.selectbox("PLATFORM", list(PLATFORMS.keys()))
        with r2c2:
            layout_opts = PLATFORMS[platform]
            layout = st.selectbox("LAYOUT / SIZE", layout_opts)
            if platform == "Other" or layout == "Custom":
                custom_size = st.text_input("Custom Size (px/cm)", label_visibility="collapsed", placeholder="Enter size...")
                if custom_size: layout = f"Custom: {custom_size}"

        # 3. Mood
        mood = st.text_input("MOOD / TONE", placeholder="e.g. Luxurious, Retro, Warm Gold")
        
        # 4. Content
        headline = st.text_input("HEADLINE (Main Text)", placeholder="Title here...")
        subtext = st.text_area("SUB-TEXT", height=68, placeholder="Secondary details...")

        # 5. Optional Details (Expander)
        with st.expander("➕ OPTIONAL INFO (Price, Contact, Notes)"):
            c_ex1, c_ex2 = st.columns(2)
            with c_ex1:
                price = st.text_input("PRICE / PROMO")
            with c_ex2:
                contact_info = st.text_input("CONTACT INFO")
            
            design_notes = st.text_area("DESIGNER NOTES", height=68)
            ref_link = st.text_input("REF LINK")

        # --- ACTIONS ---
        st.markdown("<br>", unsafe_allow_html=True)
        ac1, ac2 = st.columns(2)
        with ac1:
            if st.button("💾 SAVE TO DB", type="primary", use_container_width=True):
                if order_name:
                    summary = f"{headline[:30]}..." if headline else "Brief"
                    save_to_db(order_name, platform, deadline_str, summary)
                else:
                    st.toast("⚠️ ENTER NAME!", icon="👾")

    # --- RIGHT COLUMN: PREVIEW ---
    with col_preview:
        header_text = f"[BRIEF] - {deadline_date.strftime('%d/%m')} - {order_name if order_name else 'NAME'}"
        
        st.markdown(f"""
<div class="preview-card">
<div style="border-bottom: 2px dashed #1A237E; padding-bottom: 10px; margin-bottom: 10px; font-weight: bold; font-size: 1.2rem;">
👁️ LIVE PREVIEW
</div>
<h3 style="color: #FF6F00 !important; margin: 0;">{header_text}</h3>
<p style="margin: 0;"><b>📅 D-Line:</b> {deadline_str}</p>
<p style="margin: 0;"><b>🕹️ Specs:</b> {platform} > {layout}</p>
{f'<p style="background: #E0F7FA; display:inline-block; padding:2px 8px; border:1px solid #1A237E;">✨ {mood}</p>' if mood else ''}

<hr style="border-top: 2px dashed #1A237E; margin: 15px 0;">

{f'<h2 style="margin:0;">📢 {headline}</h2>' if headline else ''}
{f'<p style="white-space: pre-wrap;">{subtext}</p>' if subtext else ''}

{(f'<b>💰 {price}</b><br>' if price else '') + (f'<b>📞 {contact_info}</b>' if contact_info else '')}

{(f'<div style="background:#FFF3E0; padding:10px; border:1px solid #FF6F00; margin-top:10px; font-size:0.9rem;"><b>⚠️ NOTE:</b> {design_notes}</div>' if design_notes else '')}
{f'<br>🔗 <a href="{ref_link}">Reference Link</a>' if ref_link else ''}
</div>
""", unsafe_allow_html=True)

        # Generate Word Logic
        safe_name = "".join([c for c in order_name if c.isalnum() or c in (' ', '_', '-')]).strip() or "Order"
        
        def create_docx():
            doc = Document()
            doc.add_heading(header_text, 0)
            
            doc.add_heading('SPECS', level=1)
            doc.add_paragraph(f"Deadline: {deadline_str}")
            doc.add_paragraph(f"Platform: {platform} | {layout}")
            if mood: doc.add_paragraph(f"Mood: {mood}")
            
            doc.add_heading('CONTENT', level=1)
            if headline: doc.add_paragraph(f"Headline: {headline}")
            if subtext: doc.add_paragraph(f"Sub-text: {subtext}")
            if price: doc.add_paragraph(f"Price: {price}")
            if contact_info: doc.add_paragraph(f"Contact: {contact_info}")
            
            if design_notes or ref_link:
                doc.add_heading('NOTES', level=1)
                if design_notes: doc.add_paragraph(design_notes)
                if ref_link: doc.add_paragraph(f"Ref: {ref_link}")
                
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer

        with ac2: # Reuse action column from left logic? No, placed under preview in right column usually better? 
            # Or use the second half of the button row we made above?
            # User wants grid inputs on Top/Left. Preview on Right. Button placement undefined but "Primary actions" implies close proximity.
            # Let's put the download button right under the preview card.
            pass

        st.markdown("<br>", unsafe_allow_html=True)
        if order_name:
            docx_file = create_docx()
            st.download_button(
                label="📥 DOWNLOAD WORD DOC",
                data=docx_file,
                file_name=f"Brief_{safe_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                type="primary"
            )
        else:
            st.button("📥 DOWNLOAD (ENTER NAME FIRST)", disabled=True, use_container_width=True)


with tab2:
    st.markdown("### 💾 MEMORY CARD (HISTORY)")
    df = load_history()
    st.dataframe(df, use_container_width=True)
    if st.button("🔄 RELOAD DATA"):
        st.rerun()

